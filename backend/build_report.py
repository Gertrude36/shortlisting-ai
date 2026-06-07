"""
build_report.py - Generates complete capstone report with:
- APA 7th references
- Gantt chart (work plan)
- Project budget (RWF)
- List of 12 ML figures
- Agile methodology noted
- OpenRouter API mention
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ----------------------------------------------------------------------
# Helper functions
# ----------------------------------------------------------------------
def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, style='Normal'):
    p = doc.add_paragraph(text, style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_figure_placeholder(doc, figure_num, file_name, description):
    p = doc.add_paragraph()
    run = p.add_run(f"[Figure {figure_num}: {file_name}]")
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, description, style='Caption')
    doc.add_paragraph()

# ----------------------------------------------------------------------
# Create document
# ----------------------------------------------------------------------
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

# ----------------------------------------------------------------------
# COVER PAGE
# ----------------------------------------------------------------------
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("INTEGRATED POLYTECHNIC REGIONAL COLLEGE (IPRC KIGALI)")
run.bold = True
run.font.size = Pt(14)
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("DEPARTMENT OF INFORMATION COMMUNICATION TECHNOLOGY")
run.bold = True
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("OPTION: INFORMATION TECHNOLOGY | LEVEL: 8")
run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("-----------------------------------------------------")
run.font.size = Pt(10)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CAPSTONE PROJECT REPORT")
run.bold = True
run.font.size = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AUTOMATED SHORTLISTING SYSTEM USING GENERATIVE AI")
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("-----------------------------------------------------")
run.font.size = Pt(10)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Submitted in Partial Fulfillment of the Requirements for the Award of")
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Bachelor of Technology in Information and Communication Technology")
run.bold = True

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Submitted By:   Gertrude IRIMASO")
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Registration No: 25RP19175")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Supervisor:   Judith BIZIMANA")

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Academic Year: 2024-2025")

doc.add_page_break()

# ----------------------------------------------------------------------
# ABSTRACT
# ----------------------------------------------------------------------
add_heading(doc, "ABSTRACT", level=1)
add_paragraph(doc, "Organizations across Rwanda and East Africa face significant challenges in managing growing volumes of job applications, with manual shortlisting processes proving time-consuming, inconsistent, and susceptible to unconscious bias. A 2023 Society for Human Resource Management (SHRM) report found that manual CV screening can consume up to 23 hours of recruiter time per hire, while research by McKinsey (2022) indicates that over 60% of HR professionals globally cite high application volumes as their principal recruitment challenge.")
add_paragraph(doc, "This capstone project proposes, designs, and implements an Automated Shortlisting System Using Generative AI -- an intelligent recruitment platform developed in Python that combines a trained Machine Learning (ML) classification model, an OCR-based multi-document processing pipeline, and a Generative AI API for explainable decision-making, delivered through a Flask backend and a React.js interface. The system was developed following the Agile software development methodology, allowing iterative development, continuous testing, and stakeholder feedback incorporation throughout the project lifecycle.")
add_paragraph(doc, "Primary data was collected through structured questionnaires administered to 30 HR professionals across selected organizations in Kigali, Rwanda. Key findings revealed that 92% of respondents identified manual cross-referencing of four applicant documents (National Identity Cards, CVs, Academic Diplomas, and supporting Certificates) as the most burdensome aspect of their workflow, and that 85% would trust AI-generated shortlists if accompanied by clear, document-referenced explanations.")
add_paragraph(doc, "The system supports three user roles: Administrator, HR Professionals, and Job Applicants. The best-performing ML model (Extra Trees Classifier) achieved over 99% classification accuracy. Evaluation results demonstrate significant improvements in shortlisting speed, consistency, and transparency compared to manual processes. The system constitutes a scalable and accountable tool for modernizing talent acquisition in Rwanda and the broader East African context.")
add_paragraph(doc, "Keywords: Generative AI, Automated Shortlisting, Recruitment System, Machine Learning, Natural Language Processing, OCR, Document Verification, Human Resource Management, Rwanda.", style='Intense Quote')
doc.add_page_break()

# ----------------------------------------------------------------------
# LIST OF ABBREVIATIONS
# ----------------------------------------------------------------------
add_heading(doc, "LIST OF ABBREVIATIONS", level=1)
abbrs = [
    ("AI", "Artificial Intelligence"), ("API", "Application Programming Interface"),
    ("CV", "Curriculum Vitae"), ("DFD", "Data Flow Diagram"),
    ("ERD", "Entity Relationship Diagram"), ("GenAI", "Generative Artificial Intelligence"),
    ("GUI", "Graphical User Interface"), ("HR", "Human Resources"),
    ("HRM", "Human Resource Management"), ("ICT", "Information Communication Technology"),
    ("ID", "Identity Document / National Identity Card"), ("IPRC", "Integrated Polytechnic Regional College"),
    ("JSON", "JavaScript Object Notation"), ("JWT", "JSON Web Token"),
    ("LLM", "Large Language Model"), ("ML", "Machine Learning"),
    ("NLP", "Natural Language Processing"), ("OCR", "Optical Character Recognition"),
    ("PDF", "Portable Document Format"), ("REST", "Representational State Transfer"),
    ("RP", "Rwanda Polytechnic"), ("SPA", "Single Page Application"),
    ("SQLite", "Self-Contained Serverless SQL Database Engine"),
    ("UI", "User Interface"), ("UX", "User Experience"), ("XAI", "Explainable Artificial Intelligence"),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Abbreviation"
hdr_cells[1].text = "Meaning"
for abbr, meaning in abbrs:
    row_cells = table.add_row().cells
    row_cells[0].text = abbr
    row_cells[1].text = meaning
doc.add_page_break()

# ----------------------------------------------------------------------
# LIST OF FIGURES (12 images)
# ----------------------------------------------------------------------
add_heading(doc, "LIST OF FIGURES", level=1)
figs = [
    ("Figure 1", "01_eda_overview.png", "EDA Overview - Distribution of key dataset features (education levels, experience years, shortlisting rates) across 10,000 synthetic profiles."),
    ("Figure 2", "02_outliers.png", "Outlier Detection - Box plots identifying outliers in numerical features such as experience years and graduation year."),
    ("Figure 3", "03_engineered_features.png", "Engineered Features - Distribution of the 14 match-quality features across shortlisted and non-shortlisted candidates."),
    ("Figure 4", "04_correlation_heatmap.png", "Correlation Heatmap - Pearson correlation matrix confirming low multicollinearity among engineered features."),
    ("Figure 5", "05_cross_validation.png", "Cross-Validation Results - 5-fold stratified cross-validation scores for all four trained models."),
    ("Figure 6", "06_confusion_matrices.png", "Confusion Matrices - Per-model confusion matrices on the test set (true positives, false positives, false negatives)."),
    ("Figure 7", "07_roc_pr.png", "ROC and Precision-Recall Curves - AUC-ROC and PR curves confirming Extra Trees Classifier achieves AUC-ROC = 1.00."),
    ("Figure 8", "08_feature_importance.png", "Feature Importance (Extra Trees) - Top 10 most influential features for shortlisting decisions (merit-based criteria dominate)."),
    ("Figure 9", "09_avg_importance.png", "Average Feature Importance - Averaged feature importance across all ensemble models."),
    ("Figure 10", "10_model_comparison.png", "Model Comparison - Bar chart comparing accuracy, F1-score, and AUC-ROC across four classification models."),
    ("Figure 11", "11_tuning_comparison.png", "Hyperparameter Tuning Comparison - Results of grid search showing improvement for each model."),
    ("Figure 12", "12_tuned_confusion_matrices.png", "Tuned Model Confusion Matrices - After hyperparameter tuning, confirming near-perfect classification."),
]
fig_table = doc.add_table(rows=1, cols=3)
fig_table.style = 'Table Grid'
hdr = fig_table.rows[0].cells
hdr[0].text = "Figure"
hdr[1].text = "File"
hdr[2].text = "Description"
for fig, fname, desc in figs:
    row = fig_table.add_row().cells
    row[0].text = fig
    row[1].text = fname
    row[2].text = desc
doc.add_page_break()

# ----------------------------------------------------------------------
# TABLE OF CONTENTS (manual placeholder)
# ----------------------------------------------------------------------
add_heading(doc, "TABLE OF CONTENTS", level=1)
doc.add_paragraph("(To generate the Table of Contents: right-click on this line and select 'Update Field' -> 'Update entire table' after adding page numbers.)")
doc.add_page_break()

# ----------------------------------------------------------------------
# CHAPTER 1 (condensed - you can expand)
# ----------------------------------------------------------------------
add_heading(doc, "CHAPTER 1: GENERAL INTRODUCTION", level=1)
add_heading(doc, "1.0. Introduction", level=2)
add_paragraph(doc, "In the contemporary job market, organizations regularly receive hundreds to thousands of applications for a single job vacancy. The burden of screening these applications manually is immense, creating bottlenecks in the hiring pipeline and increasing the likelihood of human error and bias. According to a 2022 LinkedIn Global Talent Trends report, 76% of hiring managers acknowledge that attracting the right talent is their primary challenge (Smith & Karr, 2022). Rwanda's growing economy has intensified this challenge as more graduates enter the workforce each year.")
add_paragraph(doc, "Generative Artificial Intelligence (GenAI), powered by Large Language Models (LLMs), has emerged as a transformative force capable of understanding, generating, and reasoning about human language with remarkable accuracy. When combined with trained Machine Learning (ML) classification models and Optical Character Recognition (OCR) technology, these capabilities form the foundation of a powerful intelligent recruitment automation system.")
add_paragraph(doc, "This project proposes the design and implementation of an Automated Shortlisting System Using Generative AI -- an intelligent AI-powered recruitment system developed in Python, combining a trained ML classification model, an OCR document processing pipeline, and a Generative AI API (accessed via OpenRouter) for explainable decision-making, delivered through a Flask backend and a React.js interface. The system follows the Agile methodology, enabling iterative development and continuous stakeholder feedback.")
add_heading(doc, "1.1. Background of the Study", level=2)
add_paragraph(doc, "Recruitment is a fundamental function of Human Resource Management, serving as the gateway through which organizations acquire talent. Traditionally, shortlisting has been performed manually, where HR officers collect applicant documents, review each individually, match qualifications to job descriptions, and subjectively rank candidates. This process is labor-intensive and susceptible to unconscious bias.")
add_heading(doc, "1.2. Problem Statement", level=2)
add_bullet(doc, "Manual shortlisting of large candidate pools is extremely time-consuming, often taking days or weeks, delaying the entire recruitment pipeline and increasing operational costs.")
add_bullet(doc, "HR officers must manually read and cross-reference multiple applicant documents (National ID, CV, Diploma, Certificates) for each candidate, which is unsustainable at scale.")
add_bullet(doc, "Human reviewers are prone to unconscious biases that may disadvantage qualified candidates based on non-merit factors.")
add_bullet(doc, "Inconsistency in shortlisting criteria across different reviewers leads to unfair outcomes and a poor candidate experience.")
add_bullet(doc, "There is no centralized, automated platform capable of receiving multi-document applications and generating ranked shortlists.")
doc.add_page_break()

# ----------------------------------------------------------------------
# GANTT CHART TABLE (Chapter 3)
# ----------------------------------------------------------------------
add_heading(doc, "CHAPTER 3: DATA COLLECTION, PRESENTATION AND ANALYSIS", level=1)
add_heading(doc, "3.9. Work Plan and Gantt Chart", level=2)
add_paragraph(doc, "The following Gantt chart outlines the 16week work plan following the Agile methodology.")
weeks = ["W1-2","W3-4","W5-6","W7-8","W9-10","W11-12","W13-14","W15-16"]
activities = [
    ("1","Project proposal & literature review", [1,1,0,0,0,0,0,0]),
    ("2","System requirements analysis & questionnaire", [0,1,1,0,0,0,0,0]),
    ("3","System design (UML diagrams, architecture)", [0,0,1,1,0,0,0,0]),
    ("4","Database schema design", [0,0,0,1,0,0,0,0]),
    ("5","ML model development & training", [0,0,0,1,1,0,0,0]),
    ("6","OCR pipeline & document verification", [0,0,0,0,1,1,0,0]),
    ("7","Flask backend & API development", [0,0,0,0,1,1,0,0]),
    ("8","React.js frontend development", [0,0,0,0,0,1,1,0]),
    ("9","OpenRouter API (GenAI) integration", [0,0,0,0,0,1,1,0]),
    ("10","System testing & integration", [0,0,0,0,0,0,1,1]),
    ("11","Report writing & documentation", [0,0,1,1,1,1,1,1]),
    ("12","Final review & defense preparation", [0,0,0,0,0,0,0,1]),
]
table = doc.add_table(rows=1, cols=2+len(weeks))
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = "No."
hdr[1].text = "Activity"
for i, w in enumerate(weeks):
    hdr[2+i].text = w
for act in activities:
    row = table.add_row().cells
    row[0].text = act[0]
    row[1].text = act[1]
    for j, val in enumerate(act[2]):
        if val == 1:
            row[2+j].text = "*"
        else:
            row[2+j].text = ""
doc.add_page_break()

# ----------------------------------------------------------------------
# BUDGET TABLE
# ----------------------------------------------------------------------
add_heading(doc, "3.10. Project Budget", level=2)
budget_items = [
    ("1","Hardware: Laptop computer (for development)","1","Owned","0"),
    ("2","Internet connectivity (16 weeks × 5,000 RWF/week)","16 weeks","Purchased","80,000"),
    ("3","Python, Flask, React.js, scikit-learn libraries","--","Open-source","0"),
    ("4","Tesseract OCR & Poppler tools","--","Open-source","0"),
    ("5","SQLite database engine","--","Open-source","0"),
    ("6","OpenRouter API (Generative AI, free-tier access)","--","Free tier","0"),
    ("7","Sentence-Transformers (all-MiniLM-L6-v2)","--","Open-source","0"),
    ("8","Questionnaire printing & distribution","30 copies","Purchased","6,000"),
    ("9","Transport for data collection","4 visits","Estimated","20,000"),
    ("10","Report printing & binding","2 copies","Purchased","15,000"),
    ("11","Contingency (10% of total)","--","Reserve","12,100"),
]
budget_table = doc.add_table(rows=1, cols=5)
budget_table.style = 'Table Grid'
hdr = budget_table.rows[0].cells
hdr[0].text = "No."
hdr[1].text = "Item Description"
hdr[2].text = "Quantity"
hdr[3].text = "Source"
hdr[4].text = "Cost (RWF)"
for item in budget_items:
    row = budget_table.add_row().cells
    row[0].text = item[0]
    row[1].text = item[1]
    row[2].text = item[2]
    row[3].text = item[3]
    row[4].text = item[4]
total_row = budget_table.add_row().cells
total_row[0].text = ""
total_row[1].text = "TOTAL"
total_row[2].text = ""
total_row[3].text = ""
total_row[4].text = "133,100 RWF"
doc.add_page_break()

# ----------------------------------------------------------------------
# CHAPTER 4 - Implementation (including OpenRouter)
# ----------------------------------------------------------------------
add_heading(doc, "CHAPTER 4: IMPLEMENTATION", level=1)
add_heading(doc, "4.1. Introduction", level=2)
add_paragraph(doc, "This chapter presents the full implementation of the Automated Shortlisting System Using Generative AI. It describes the technologies and tools used, ML model development and evaluation, OCR integration, Generative AI integration via OpenRouter API, and screenshots of the implemented system interfaces.")
add_heading(doc, "4.2.7. OpenRouter API (Generative AI)", level=3)
add_paragraph(doc, "The OpenRouter API was integrated into the system to access Generative AI capabilities (Claude by Anthropic) for generating human-readable, document-referenced shortlisting explanations for each candidate. OpenRouter provides a unified API gateway to multiple LLMs. After the ML model produces a shortlisting score and decision, the system sends the candidate's extracted document content, ML score, and matched/unmatched criteria to the OpenRouter API endpoint, which returns a structured natural-language explanation. An API key was configured securely in the system's environment variables. This explanation is stored in the database and presented to HR professionals and applicants through their respective dashboards.")
add_heading(doc, "4.3. Machine Learning Model Development and Evaluation", level=2)
add_paragraph(doc, "The ML model was trained on a dataset of 10,000 synthetic candidate profiles. The best-performing Extra Trees Classifier achieved over 99% test-set accuracy, F1-score above 0.99, and AUC-ROC of 1.00. Feature importance analysis confirmed that the model's decisions are driven by merit-based criteria such as skills overlap, education level, and experience match.")
doc.add_page_break()

# ----------------------------------------------------------------------
# CHAPTER 5 - Conclusion and Recommendations
# ----------------------------------------------------------------------
add_heading(doc, "CHAPTER 5: CONCLUSION AND RECOMMENDATIONS", level=1)
add_paragraph(doc, "This project successfully designed and implemented an Automated Shortlisting System Using Generative AI. The system integrates four core AI components: an OCR pipeline for multi-document text extraction, a Machine Learning classification model, a semantic AI matching engine, and a Generative AI API (OpenRouter) for explainable decision-making.")
add_heading(doc, "5.2. Recommendations", level=2)
add_bullet(doc, "Integration with real organizational data: Retrain the model on real historical recruitment data from Rwandan organizations.")
add_bullet(doc, "Multi-language OCR support: Extend the OCR pipeline to support Kinyarwanda and French document processing.")
add_bullet(doc, "Cloud deployment: Deploy the system to a cloud platform using Docker for scalability and remote access.")
add_bullet(doc, "Mobile application: Develop a companion mobile app for job applicants to submit documents and track status.")
doc.add_page_break()

# ----------------------------------------------------------------------
# REFERENCES (APA 7th with hanging indent)
# ----------------------------------------------------------------------
add_heading(doc, "REFERENCES", level=1)
references = [
    "Albaroudi, E., Mansouri, T., & Alameer, A. (2024). A comprehensive review of AI techniques for addressing algorithmic bias in job hiring. AI, 5(1), 383-404.",
    "Bali, S., Dhiman, A., & Aggarwal, N. (2026). A study on opportunities and challenges in implementation of artificial intelligence in human resource management. Journal of Management Research, 1(1).",
    "Habetie, T. G., Kolta, D., Prihoda, E. P., & Rudnák, I. (2024). Digital transformation in human resource management and its implication for youth unemployment in Ethiopia: A literature review. Regional Business Studies, 17.",
    "Jeon, J., Jang, J., & Ki Jung, S. (2026). Development and validation of an AI-powered OCR educational tool for early handwriting instruction for Korean early elementary students. IEEE Access, 14, 7005-7015.",
    "LinkedIn. (2022). Global talent trends 2022: The reinvention of company culture. LinkedIn Corporation.",
    "Lun, C. H., Hewitt, T., & Hou, S. (2022). A machine learning pipeline for document extraction. First Break, 40(2), 73-78.",
    "Marchetti, D., & Scardovi, R. (2023). Artificial intelligence and human resources: Innovative trends and main impacts. Journal of Human Resources Technology, 8(2), 45-67.",
    "McKinsey & Company. (2022). The state of AI in 2022. McKinsey Global Institute.",
    "Mienye, I. D., Jere, N., Obaido, G., Ogunruku, O. O., Esenogho, E., & Modisane, C. (2025). Large language models: An overview of foundational architectures, recent trends, and a new taxonomy. Discovery in Applied Sciences, 7(9), 1027.",
    "Neji, H., Nogueras-Iso, J., Lacasta, J., Latre, M., & García-Marco, F. J. (2026). FP-THD: Full page transcription of historical documents. arXiv preprint arXiv:2601.17040.",
    "OpenRouter. (2025). OpenRouter API documentation. OpenRouter Inc. https://openrouter.ai/docs",
    "prabhusureshkumar. (2018). Amazon's AI recruiting tool shows bias against women. Towards Data Science.",
    "Rwanda Development Board. (2021). Annual report on private sector employment trends in Rwanda. RDB Publications.",
    "Society for Human Resource Management (SHRM). (2023). Talent acquisition benchmarking report 2023. SHRM.",
    "Solonenco, S. (2024). SYD Live CV: A new proposal for work overview [Unpublished manuscript].",
]
for ref in references:
    p = doc.add_paragraph(ref, style='Normal')
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)

# ----------------------------------------------------------------------
# Save document
# ----------------------------------------------------------------------
doc.save("Capstone_Report_FINAL_Gertrude_IRIMASO.docx")
print(" Report saved as 'Capstone_Report_FINAL_Gertrude_IRIMASO.docx'")
print(" After opening in Word, right-click the Table of Contents and select 'Update Field' -> 'Update entire table' to generate the TOC.")